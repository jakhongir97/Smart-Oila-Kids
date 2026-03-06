from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from datetime import date

OUT_PATH = "output/doc/smart_oila_kids_6_week_roadmap.docx"


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def main():
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)

    title = doc.add_heading("Smart Oila Kids - 6-Week Roadmap (Repo-Based)", level=0)
    title.alignment = 0

    doc.add_paragraph("Generated: March 5, 2026")
    doc.add_paragraph("Planning window: Week 1 (March 9-15, 2026) through Week 6 (April 13-19, 2026)")
    doc.add_paragraph("Source basis: current Swift modules, service wiring, CHILD_ONLY_EXTRACTION_SPEC.md, PARENT_CHILD_TESTING.md, and scripts/check_openapi_coverage.py output.")

    doc.add_heading("1) Current Repo Baseline", level=1)
    add_bullets(doc, [
        "Auth is implemented with QR claim + legacy fallback claim endpoint (/upload-v2/device), plus child binding verification.",
        "Root flow is DSN-gated: missing DSN -> Auth, existing DSN -> Main.",
        "Main includes SOS, weekly usage chart, device status snapshot, unread chat/tasks/notifications counters, and pull-to-refresh.",
        "Chat includes history fetch, message send (multipart), websocket receive, local outbox queue, and read-state tracking.",
        "Tasks include awards fetch, change-status action, optimistic queueing for retry, and local cache fallback.",
        "Settings include profile rename, connected device rename/avatar upload/delete, and local cache.",
        "Background services include geo websocket (location + system_info events), device lock polling/overlay, and push inbox/deep-link routing.",
        "App has localization in en/ru/uz and a shared design system (colors, typography, child chrome components).",
    ])

    doc.add_heading("2) Architecture Constraints and Repo-TODO Signals", level=1)
    add_bullets(doc, [
        "Platform constraints: SwiftUI app, iOS deployment target 16.0, single iPhone family target.",
        "Scope constraint: child-only app boundaries are documented; parent-only feature families are out of scope.",
        "API/WS constraints: env-driven base URLs with hardcoded defaults and websocket secret path tokening.",
        "Auth constraint: mixed DSN-only and tokenized endpoint behavior is actively handled (401/403/404 fallbacks).",
        "Coverage gap signal: child currently covers 19/85 REST operations and 2/23 websocket routes from OpenAPI coverage audit (run on March 5, 2026).",
        "Integration dependency: parent-child simulator/testing scripts depend on sibling parent repo path on Desktop.",
        "Testing constraint: automated coverage is minimal (script-level unit test only for OpenAPI coverage parser).",
        "Backlog/TODO signal from extraction spec: confirm official child HTTPS base URL, auth/token rules, websocket routes, and finalize unresolved Settings node mapping.",
    ])

    doc.add_heading("3) Milestones", level=1)
    add_bullets(doc, [
        "Milestone 1 (Week 1): Scope and API contract freeze for child app parity.",
        "Milestone 2 (Week 2): Auth and session reliability hardening.",
        "Milestone 3 (Week 3): Core engagement reliability (Main + Chat + Tasks).",
        "Milestone 4 (Week 4): Background services and lock/push correctness.",
        "Milestone 5 (Week 5): Settings/device management and localization polish.",
        "Milestone 6 (Week 6): Release candidate stabilization and pilot readiness.",
    ])

    doc.add_heading("4) Week-by-Week Plan", level=1)
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = ["Week", "Focus", "Weekly Goals", "Deliverables", "Dependencies", "Risks"]
    for i, text in enumerate(headers):
        table.rows[0].cells[i].text = text

    rows = [
        [
            "Week 1\nMar 9-15",
            "Scope + Contract Lock",
            "Lock child-only scope and parity baseline.\nRun endpoint audit + OpenAPI gap baseline.\nResolve Settings mapping ambiguity.",
            "Approved scope matrix (feature in/out).\nContract matrix for Auth/Chat/Tasks/Geo/Lock routes.\nPrioritized 6-week engineering backlog.",
            "Backend owner confirmations.\nFigma access for unresolved Settings node.",
            "Risk: contract drift during build.\nMitigation: freeze doc + change control.",
        ],
        [
            "Week 2\nMar 16-22",
            "Auth + Session Hardening",
            "Harden QR and legacy claim path behavior.\nValidate DSN verification retries and error states.\nStabilize token refresh/session persistence edge cases.",
            "Auth reliability checklist pass.\nKnown error-state UX matrix and expected copy.\nSession-handling regression checklist.",
            "Valid QR payload corpus.\nTest credentials for tokenized APIs.",
            "Risk: 401/404 loops and broken onboarding.\nMitigation: explicit retry caps + fallback gates.",
        ],
        [
            "Week 3\nMar 23-29",
            "Main + Chat + Tasks Reliability",
            "Improve websocket reconnect/outbox replay behavior.\nHarden unread counters and cross-screen sync.\nValidate queued task actions and cache replay.",
            "Chat reliability report (delivery/reconnect).\nTasks offline-sync validation report.\nMain dashboard fallback behavior verified.",
            "Stable websocket environment.\nPush payload samples for chat/tasks events.",
            "Risk: duplicate/out-of-order messages.\nMitigation: timestamp/fingerprint guardrails + replay tests.",
        ],
        [
            "Week 4\nMar 30-Apr 5",
            "Background + Lock + Push",
            "Validate geo websocket payload cadence in foreground/background.\nVerify lock overlay correctness from full/global lock endpoints.\nTest notification deep-links and badge reconciliation.",
            "Background behavior matrix (device/network conditions).\nLock-state acceptance checks.\nPush inbox/deep-link end-to-end test evidence.",
            "APNs test setup.\nReal device coverage beyond simulators.",
            "Risk: iOS background delivery constraints.\nMitigation: fallback polling + monitoring hooks.",
        ],
        [
            "Week 5\nApr 6-12",
            "Settings + Polish",
            "Finalize profile/device rename/avatar/delete flows.\nClose localization and UI consistency gaps.\nRun regression pass on settings and session clear flows.",
            "Settings regression suite results.\nLocalization QA sheet (en/ru/uz).\nDesign and copy sign-off package.",
            "Backend device-management permissions.\nFinal UX approvals.",
            "Risk: permission mismatches on device mutation endpoints.\nMitigation: preflight authorization checks + safe fallbacks.",
        ],
        [
            "Week 6\nApr 13-19",
            "Release Readiness",
            "Run parent-child integration scripts and smoke suite.\nRe-run OpenAPI coverage and verify targeted closure.\nStabilize, cut RC, and prepare pilot monitoring.",
            "Release candidate build.\nGo/no-go checklist with rollback plan.\nPilot rollout + post-launch monitoring checklist.",
            "Parent repo availability for integration run.\nRelease sign-off from backend and product.",
            "Risk: late backend changes or integration regressions.\nMitigation: freeze window + rollback criteria.",
        ],
    ]

    for row_data in rows:
        row = table.add_row().cells
        for idx, value in enumerate(row_data):
            row[idx].text = value

    doc.add_heading("5) Explicit Dependencies", level=1)
    add_bullets(doc, [
        "Backend API owner sign-off on child auth/message/awards/lock/geo contracts.",
        "Websocket environment stability (base URL, secret path, and access policy).",
        "Access to parent repository path used by integration scripts (../Smart Oila Parent).",
        "Figma access to finalize unresolved Settings parity mapping.",
        "QA devices and APNs-capable environment for push and background validation.",
    ])

    doc.add_heading("6) Explicit Risk Register", level=1)
    add_bullets(doc, [
        "Contract drift risk between legacy and tokenized endpoints while migration is active.",
        "Authentication mode mismatch risk (DSN-only behavior vs token-required routes).",
        "Websocket resilience risk under reconnect/load/network changes.",
        "Background execution risk on iOS for timely location/system_info delivery.",
        "Low automated test coverage risk causing regressions in multi-feature interactions.",
    ])

    doc.add_heading("7) Definition of Done by End of Week 6", level=1)
    add_bullets(doc, [
        "Core child flows are reliable in production-like conditions: Auth, Main, Chat, Tasks, Settings, Geo, Lock, Push.",
        "Critical endpoint and websocket paths for child use-cases are validated against OpenAPI and runtime behavior.",
        "Integration scripts and smoke checks pass for parent-child workflow on target simulators/devices.",
        "Release candidate and pilot rollout checklist are complete with clear rollback criteria.",
    ])

    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    main()
